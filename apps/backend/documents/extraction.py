"""AILIZA -- einheitliche, strukturierte Dokument-Extraktion.

Ein gemeinsames Modul fuer alle Formate (PDF/DOCX/XLSX/CSV, optional Bilder
per lokalem OCR), das bisher an zwei getrennten Stellen (documents/
document_handler.py und knowledge/ingestion.py) jeweils eigenstaendig und
nur oberflaechlich (reiner Fliesstext, keine Seiten-/Abschnittsinformation)
implementiert war.

WICHTIG (Karo-Entscheidung 2026-08-03):
- Bilder/OCR sind bewusst NUR lokal verfuegbar (System-Abhaengigkeit
  tesseract-ocr laesst sich auf dem aktuellen Server-Hosting/Render nicht
  ohne Weiteres installieren). Ob OCR verfuegbar ist, wird zur Laufzeit an
  zwei Bedingungen geprueft: (1) die Python-Pakete pytesseract/Pillow sind
  installiert, UND (2) AILIZA_LOCAL_OCR_ENABLED=true ist gesetzt. Fehlt
  eine der beiden Bedingungen, gelten Bild-Dateitypen als NICHT unterstuetzt
  (fail-closed, kein stiller Leerlauf).
- Fehlt eine Bibliothek (z.B. pdfplumber nicht installiert), wird das NICHT
  mehr stillschweigend als leerer Text zurueckgegeben (das war ein
  bestehendes Zertifizierungsrisiko: unklassifizierter Inhalt waere sonst
  faelschlich als "gepasst" durchgewunken worden). Stattdessen liefert
  ExtractionResult ein explizites `warnings`-Signal, das der Aufrufer
  auswerten MUSS (document_handler.py/knowledge/ingestion.py blockieren in
  diesem Fall den Upload mit einer verstaendlichen Meldung).
"""
from __future__ import annotations

import csv
import io
import os
from dataclasses import dataclass, field

# Sicherheitslimits: begrenzen die Verarbeitungszeit/den Speicherbedarf
# durch harte Obergrenzen auf die Menge der verarbeiteten Elemente (Seiten/
# Zeilen/Zellen), statt eines Prozess-Timeouts -- das ist deterministisch
# und plattformunabhaengig (kein signal.alarm, funktioniert auch unter
# Windows fuer die geplante Desktop-Distribution, siehe Roadmap-Doku).
MAX_PDF_PAGES = 200
MAX_DOCX_PARAGRAPHS = 20_000
MAX_DOCX_TABLE_ROWS = 5_000
MAX_XLSX_CELLS_PER_SHEET = 500_000
MAX_XLSX_SHEETS = 50
MAX_CSV_ROWS = 50_000
MAX_IMAGE_PIXELS = 40_000_000  # gegen Pillow-Decompression-Bomb

_OCR_ENV_FLAG = "AILIZA_LOCAL_OCR_ENABLED"


@dataclass
class ExtractedSection:
    page_number: int | None
    section_title: str | None
    content_markdown: str
    element_type: str  # "paragraph" | "heading" | "table" | "ocr_text" | "row"


@dataclass
class ExtractionResult:
    sections: list[ExtractedSection] = field(default_factory=list)
    full_text: str = ""
    page_count: int | None = None
    truncated: bool = False
    warnings: list[str] = field(default_factory=list)

    @property
    def library_missing(self) -> bool:
        return any(w.startswith("library_missing:") for w in self.warnings)


def ocr_available() -> bool:
    """OCR gilt nur als verfuegbar, wenn die Pakete installiert sind UND
    die lokale Freigabe (AILIZA_LOCAL_OCR_ENABLED=true) gesetzt ist."""
    if os.getenv(_OCR_ENV_FLAG, "").strip().lower() not in {"1", "true", "yes"}:
        return False
    try:
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
    except ImportError:
        return False
    return True


def _extract_pdf(content: bytes) -> ExtractionResult:
    try:
        import pdfplumber
    except ImportError:
        return ExtractionResult(warnings=["library_missing:pdfplumber"])

    result = ExtractionResult()
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            result.page_count = len(pdf.pages)
            for page_index, page in enumerate(pdf.pages):
                if page_index >= MAX_PDF_PAGES:
                    result.truncated = True
                    result.warnings.append(f"truncated:pdf_pages>{MAX_PDF_PAGES}")
                    break
                text = page.extract_text() or ""
                if text.strip():
                    result.sections.append(ExtractedSection(
                        page_number=page_index + 1, section_title=None,
                        content_markdown=text, element_type="paragraph",
                    ))
                for table in page.extract_tables() or []:
                    md = _table_to_markdown(table)
                    if md:
                        result.sections.append(ExtractedSection(
                            page_number=page_index + 1, section_title=None,
                            content_markdown=md, element_type="table",
                        ))
    except Exception as exc:  # defekte/verschluesselte PDF -- fail-closed, kein Absturz
        result.warnings.append(f"extraction_error:{type(exc).__name__}")

    result.full_text = "\n\n".join(s.content_markdown for s in result.sections)
    return result


def _extract_docx(content: bytes) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError:
        return ExtractionResult(warnings=["library_missing:python-docx"])

    result = ExtractionResult()
    try:
        doc = Document(io.BytesIO(content))
        for index, paragraph in enumerate(doc.paragraphs):
            if index >= MAX_DOCX_PARAGRAPHS:
                result.truncated = True
                result.warnings.append(f"truncated:docx_paragraphs>{MAX_DOCX_PARAGRAPHS}")
                break
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name if paragraph.style else "") or ""
            is_heading = style_name.lower().startswith("heading") or style_name.lower() == "title"
            result.sections.append(ExtractedSection(
                page_number=None,
                section_title=text if is_heading else None,
                content_markdown=text, element_type="heading" if is_heading else "paragraph",
            ))
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows[:MAX_DOCX_TABLE_ROWS]]
            md = _table_to_markdown(rows)
            if md:
                result.sections.append(ExtractedSection(
                    page_number=None, section_title=None,
                    content_markdown=md, element_type="table",
                ))
    except Exception as exc:
        result.warnings.append(f"extraction_error:{type(exc).__name__}")

    result.full_text = "\n\n".join(s.content_markdown for s in result.sections)
    return result


def _extract_xlsx(content: bytes) -> ExtractionResult:
    try:
        import openpyxl
    except ImportError:
        return ExtractionResult(warnings=["library_missing:openpyxl"])

    result = ExtractionResult()
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        for sheet_index, ws in enumerate(wb.worksheets):
            if sheet_index >= MAX_XLSX_SHEETS:
                result.truncated = True
                result.warnings.append(f"truncated:xlsx_sheets>{MAX_XLSX_SHEETS}")
                break
            rows: list[list[str]] = []
            cell_count = 0
            for row in ws.iter_rows(values_only=True):
                cell_count += len(row)
                rows.append(["" if c is None else str(c) for c in row])
                if cell_count >= MAX_XLSX_CELLS_PER_SHEET:
                    result.truncated = True
                    result.warnings.append(f"truncated:xlsx_cells>{MAX_XLSX_CELLS_PER_SHEET}")
                    break
            md = _table_to_markdown(rows)
            if md:
                result.sections.append(ExtractedSection(
                    page_number=None, section_title=ws.title,
                    content_markdown=md, element_type="table",
                ))
    except Exception as exc:
        result.warnings.append(f"extraction_error:{type(exc).__name__}")

    result.full_text = "\n\n".join(s.content_markdown for s in result.sections)
    return result


def _extract_csv(content: bytes) -> ExtractionResult:
    result = ExtractionResult()
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        result.warnings.append(f"extraction_error:{type(exc).__name__}")
        return result

    try:
        reader = csv.reader(io.StringIO(text))
        rows: list[list[str]] = []
        for index, row in enumerate(reader):
            if index >= MAX_CSV_ROWS:
                result.truncated = True
                result.warnings.append(f"truncated:csv_rows>{MAX_CSV_ROWS}")
                break
            rows.append(row)
    except csv.Error as exc:
        result.warnings.append(f"extraction_error:{type(exc).__name__}")
        return result

    md = _table_to_markdown(rows)
    if md:
        result.sections.append(ExtractedSection(
            page_number=None, section_title=None, content_markdown=md, element_type="table",
        ))
    result.full_text = md
    return result


def _extract_image(content: bytes) -> ExtractionResult:
    if not ocr_available():
        return ExtractionResult(warnings=["library_missing:ocr_local_only"])

    import pytesseract
    from PIL import Image

    result = ExtractionResult()
    try:
        Image.MAX_IMAGE_PIXELS = MAX_IMAGE_PIXELS
        image = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(image, lang="deu+eng")
    except Exception as exc:
        result.warnings.append(f"extraction_error:{type(exc).__name__}")
        return result

    text = text.strip()
    if text:
        result.sections.append(ExtractedSection(
            page_number=None, section_title=None, content_markdown=text, element_type="ocr_text",
        ))
    result.full_text = text
    return result


def _table_to_markdown(rows: list[list[str]]) -> str:
    """Baut eine simple Markdown-Tabelle. Leere Tabellen liefern "" (werden
    vom Aufrufer nicht als eigene Section aufgenommen)."""
    rows = [r for r in rows if any(str(c).strip() for c in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    padded = [list(r) + [""] * (width - len(r)) for r in rows]

    def _fmt_row(r: list[str]) -> str:
        return "| " + " | ".join(str(c).replace("\n", " ").replace("|", "\\|") for c in r) + " |"

    lines = [_fmt_row(padded[0]), "| " + " | ".join(["---"] * width) + " |"]
    lines.extend(_fmt_row(r) for r in padded[1:])
    return "\n".join(lines)


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".csv": _extract_csv,
}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}


def supported_structured_extensions() -> set[str]:
    """Alle Erweiterungen, fuer die dieses Modul strukturierte Extraktion
    anbietet -- Bild-Formate nur, wenn OCR zur Laufzeit tatsaechlich
    verfuegbar ist (siehe ocr_available())."""
    extensions = set(_EXTRACTORS)
    if ocr_available():
        extensions |= _IMAGE_EXTENSIONS
    return extensions


def extract_structured(ext: str, content: bytes) -> ExtractionResult:
    """Zentrale Extraktions-Funktion fuer alle nicht-reinen Textformate.
    Reines TXT/MD wird bewusst NICHT hier behandelt -- dafuer reicht ein
    einfacher UTF-8-Decode, der beim Aufrufer bleibt (kein Mehrwert durch
    Struktur-Extraktion bei bereits reinem Text)."""
    ext = ext.lower()
    if ext in _IMAGE_EXTENSIONS:
        return _extract_image(content)
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        return ExtractionResult(warnings=[f"unsupported_extension:{ext}"])
    return extractor(content)
